from docx import Document
from pypdf import PdfReader
# from pdfminer.high_level import extract_text as fallback_text_extraction


def read_pdf(file_path):
    text = ""
    try:
        reader = PdfReader(file_path)
        for page in reader.pages:
            text += page.extract_text()
    except Exception as exc:
        text = 'error is {}'.format(exc)
    return text


def read_docx(file_path):
    doc = Document(file_path)
    content = ''
    # 每一段的编号、内容
    for i in range(len(doc.paragraphs)):
        istr = str(i) + " " + doc.paragraphs[i].text + "\n"
        content += istr
    # 表格
    tbs = doc.tables
    for tb in tbs:
        # 行
        for row in tb.rows:
            # 列
            for cell in row.cells:
                cell_str = cell.text + "\t"
                content += cell_str
            content += '\n'
    return content


def read_file(file_path: str):
    if file_path.endswith('.pdf'):
        return read_pdf(file_path)
    elif file_path.endswith('.doc') or file_path.endswith('.docx'):
        return read_docx(file_path)
    else:
        return ''


if __name__ == '__main__':
    content = read_pdf('../static/pdf_and_doc/刘宁-算法工程师.pdf')
    print(content)
